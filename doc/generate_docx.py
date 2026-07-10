#!/usr/bin/env python3
"""
OfflineAIStudio 技术设计文档 -> DOCX 转换脚本
将 Markdown 文档转为专业排版的 Word 文档，Mermaid 图表通过 mermaid.ink 渲染为 PNG 嵌入。
"""

import base64
import json
import os
import re
import sys
import urllib.request
import urllib.error
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 配置
# ============================================================
SCRIPT_DIR = Path(__file__).parent
MD_FILE = SCRIPT_DIR / "技术设计文档.md"
OUTPUT_DOCX = SCRIPT_DIR / "OfflineAIStudio_技术设计文档.docx"
IMG_DIR = SCRIPT_DIR / "_diagrams"
DIAGRAM_DPI = 150  # 图片宽度（英寸）

# ============================================================
# Mermaid -> PNG 渲染
# ============================================================

def mermaid_to_png_url(mermaid_code: str) -> str:
    """不再使用 URL 方式，改用 POST 请求"""
    return ""


def download_image(url: str, save_path: Path) -> bool:
    """下载图片到本地（保留兼容）"""
    try:
        req = urllib.request.Request(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) OfflineAIStudio-DocGen'
        })
        with urllib.request.urlopen(req, timeout=30) as resp:
            if resp.status == 200:
                with open(save_path, 'wb') as f:
                    f.write(resp.read())
                return True
    except Exception as e:
        print(f"  [WARN] 下载失败: {e}")
    return False


def render_mermaid_via_post(mermaid_code: str, save_path: Path) -> bool:
    """通过 POST 请求 kroki.io 渲染 Mermaid 图表为高清 PNG"""
    try:
        # 使用 scale=2 提高分辨率
        payload = json.dumps({
            "diagram_source": mermaid_code,
            "diagram_type": "mermaid",
            "output_format": "png",
            "options": {
                "scale": 2,  # 2倍分辨率
                "theme": "default"
            }
        }).encode('utf-8')

        req = urllib.request.Request(
            "https://kroki.io/mermaid/png",
            data=payload,
            headers={
                'Content-Type': 'application/json',
                'Accept': 'image/png',
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36',
                'Accept-Language': 'en-US,en;q=0.9',
                'Origin': 'https://kroki.io',
                'Referer': 'https://kroki.io/',
            },
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = resp.read()
            # 检查是否是有效的 PNG 文件（以 PNG magic bytes 开头）
            if len(data) > 2000 and data[:8] == b'\x89PNG\r\n\x1a\n':
                with open(save_path, 'wb') as f:
                    f.write(data)
                return True
            elif len(data) > 1000:
                with open(save_path, 'wb') as f:
                    f.write(data)
                return True
            else:
                print(f"  [WARN] 响应数据过小 ({len(data)} bytes)")
    except urllib.error.HTTPError as e:
        # 有些情况下返回了图片但 HTTP 状态码非 200，尝试读取响应体
        body = e.read()
        if len(body) > 2000 and body[:8] == b'\x89PNG\r\n\x1a\n':
            with open(save_path, 'wb') as f:
                f.write(body)
            print(f"  [WARN] HTTP {e.code} 但收到有效高清 PNG")
            return True
        print(f"  [WARN] HTTP {e.code}: {body[:80]}")
    except Exception as e:
        print(f"  [WARN] 请求失败: {e}")
    return False


def render_mermaid_diagram(index: int, mermaid_code: str) -> Path | None:
    """渲染单个 Mermaid 图表，返回图片路径"""
    IMG_DIR.mkdir(exist_ok=True)
    img_path = IMG_DIR / f"diagram_{index:02d}.png"

    if img_path.exists() and img_path.stat().st_size > 500:
        print(f"  [SKIP] {img_path.name} 已存在")
        return img_path

    print(f"  [RENDER] 图表 {index} via POST to kroki.io...")
    if render_mermaid_via_post(mermaid_code, img_path):
        return img_path

    # 清理失败文件
    if img_path.exists():
        img_path.unlink()
    return None


# ============================================================
# Markdown 解析
# ============================================================

def extract_mermaid_blocks(md_text: str) -> list[tuple[str, str]]:
    """提取所有 Mermaid 代码块，返回 [(类型, 代码)]"""
    pattern = r'```mermaid\s*\n(.*?)```'
    matches = re.findall(pattern, md_text, re.DOTALL)
    results = []
    for code in matches:
        first_line = code.strip().split('\n')[0].strip()
        diagram_type = "diagram"
        for kw in ['flowchart', 'sequenceDiagram', 'classDiagram', 'stateDiagram',
                    'erDiagram', 'gantt', 'pie', 'mindmap', 'gitGraph', 'journey']:
            if kw in first_line:
                diagram_type = kw
                break
        results.append((diagram_type, code.strip()))
    return results


# ============================================================
# DOCX 文档生成
# ============================================================

class DocxBuilder:
    """构建专业排版的 Word 文档"""

    # 颜色常量
    COLOR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)      # 靛蓝标题
    COLOR_SECONDARY = RGBColor(0x37, 0x41, 0x51)     # 深灰副标题
    COLOR_ACCENT = RGBColor(0x4F, 0x46, 0xE5)        # 紫色强调
    COLOR_BODY = RGBColor(0x1F, 0x29, 0x37)          # 正文深灰
    COLOR_TABLE_HEADER_BG = "1E3A8A"
    COLOR_TABLE_ALT_BG = "EFF6FF"
    COLOR_BORDER = RGBColor(0xBF, 0xDB, 0xFE)

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_page(self):
        """页面设置"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _setup_styles(self):
        """配置文档样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        font.color.rgb = self.COLOR_BODY
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        # 标题样式
        for level, (size, color) in enumerate([
            (Pt(26), self.COLOR_PRIMARY),    # Heading 1
            (Pt(20), self.COLOR_PRIMARY),    # Heading 2
            (Pt(15), self.COLOR_ACCENT),     # Heading 3
            (Pt(13), self.COLOR_SECONDARY),  # Heading 4
        ], 1):
            hs = self.doc.styles[f'Heading {level}']
            hs.font.name = 'Microsoft YaHei'
            hs.font.size = size
            hs.font.color.rgb = color
            hs.font.bold = True
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            hs.paragraph_format.space_after = Pt(8)

    def add_cover_page(self):
        """添加封面页"""
        # 空行填充
        for _ in range(6):
            self.doc.add_paragraph('')

        # 标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('OfflineAIStudio')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_PRIMARY
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 副标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('技术设计文档')
        run.font.size = Pt(28)
        run.font.color.rgb = self.COLOR_ACCENT
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        self.doc.add_paragraph('')

        # 分隔线
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 40)
        run.font.color.rgb = self.COLOR_BORDER
        run.font.size = Pt(14)

        self.doc.add_paragraph('')

        # 文档信息表格
        info_data = [
            ('文档版本', 'v1.0'),
            ('项目版本', 'v2.0.0'),
            ('技术栈', 'C++20 / Qt 6.11+ / Multi-Agent Architecture'),
            ('编写日期', '2026年7月'),
            ('文档密级', '内部公开'),
        ]
        table = self.doc.add_table(rows=len(info_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for i, (key, val) in enumerate(info_data):
            cell_k = table.cell(i, 0)
            cell_v = table.cell(i, 1)
            cell_k.text = key
            cell_v.text = val
            for cell in [cell_k, cell_v]:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Microsoft YaHei'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            # 键列加粗加背景
            for para in cell_k.paragraphs:
                for run in para.runs:
                    run.font.bold = True
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.COLOR_TABLE_HEADER_BG}"/>')
            cell_k.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell_k._tc.get_or_add_tcPr().append(shading)

        self.doc.add_page_break()

    def add_toc_page(self, sections: list[str]):
        """添加目录页"""
        self.doc.add_heading('目录', level=1)
        for i, title in enumerate(sections, 1):
            p = self.doc.add_paragraph()
            run = p.add_run(f'{i}.  {title}')
            run.font.size = Pt(12)
            run.font.color.rgb = self.COLOR_SECONDARY
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                      font_size: float | None = None, color: RGBColor | None = None,
                      alignment=None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
        if alignment is not None:
            p.alignment = alignment
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        return p

    def add_rich_paragraph(self, segments: list[tuple[str, dict]]):
        """添加富文本段落。segments: [(text, {bold, italic, color, size})]"""
        p = self.doc.add_paragraph()
        for text, opts in segments:
            run = p.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if opts.get('bold'):
                run.font.bold = True
            if opts.get('italic'):
                run.font.italic = True
            if opts.get('size'):
                run.font.size = Pt(opts['size'])
            if opts.get('color'):
                run.font.color.rgb = opts['color']
        return p

    def add_bullet(self, text: str, level: int = 0):
        p = self.doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
        return p

    def add_table(self, headers: list[str], rows: list[list[str]],
                  col_widths: list[float] | None = None):
        """添加格式化表格"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = header
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.COLOR_TABLE_HEADER_BG}"/>')
            cell._tc.get_or_add_tcPr().append(shading)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 数据行
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.cell(i + 1, j)
                cell.text = val
                # 交替行背景
                if i % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.COLOR_TABLE_ALT_BG}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9.5)
                        run.font.name = 'Microsoft YaHei'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 列宽
        if col_widths:
            for j, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[j].width = Cm(width)

        self.doc.add_paragraph('')  # 表后空行
        return table

    def add_image(self, img_path: str | Path, width: float = 6.5, caption: str = '', 
               description: str = ''):
        """添加居中图片及描述"""
        # 图片
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width))
        
        # 图片说明（斜体、居中）
        if caption:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(caption)
            run2.font.size = Pt(10)
            run2.font.bold = True
            run2.font.color.rgb = self.COLOR_SECONDARY
            run2.font.name = 'Microsoft YaHei'
            run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 图片详细描述
        if description:
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p3.paragraph_format.first_line_indent = Cm(0.75)
            run3 = p3.add_run(description)
            run3.font.size = Pt(10)
            run3.font.italic = True
            run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run3.font.name = 'Microsoft YaHei'
            run3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    def add_code_block(self, code: str, language: str = ''):
        """添加代码块（带灰色背景）"""
        p = self.doc.add_paragraph()
        # 添加灰色背景的边框效果
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>')
        pPr.append(shading)
        # 添加边框
        borders = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="94A3B8"/>'
            f'</w:pBdr>'
        )
        pPr.append(borders)

        run = p.add_run(code)
        run.font.size = Pt(8.5)
        run.font.name = 'Consolas'
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, path: str | Path):
        self.doc.save(str(path))


# ============================================================
# 主文档构建
# ============================================================

def build_document(diagrams: dict[int, Path]):
    """构建完整文档"""
    builder = DocxBuilder()
    D = diagrams  # 缩写

    # ==================== 封面 ====================
    builder.add_cover_page()

    # ==================== 目录 ====================
    toc_sections = [
        '项目概述',
        '技术架构总览',
        '核心系统设计',
        '关键业务流程',
        '数据模型设计',
        '接口设计',
        '状态机设计',
        '安全架构设计',
        '部署与运行架构',
        '技术选型与约束',
        '扩展性设计',
        '项目演进路线',
    ]
    builder.add_toc_page(toc_sections)

    # ==================== 第一章：项目概述 ====================
    builder.add_heading('一、项目概述', level=1)

    builder.add_heading('1.1 项目背景', level=2)
    builder.add_paragraph(
        '随着大语言模型（LLM）技术的快速发展，企业对于数据隐私和离线可控的智能助手需求日益增强。'
        '市面上主流方案采用"大模型直连执行"模式，即 LLM 直接决定调用哪些工具、执行哪些命令，存在以下风险：'
    )
    builder.add_bullet('安全风险：大模型可能产生"幻觉"，执行危险或不可预期的系统操作')
    builder.add_bullet('可控性弱：执行过程黑盒化，无法审计、无法中断')
    builder.add_bullet('离线能力缺失：依赖云端 API 和网络连接，无法在内网/隔离环境运行')

    builder.add_heading('1.2 项目定位', level=2)
    builder.add_rich_paragraph([
        ('OfflineAIStudio', {'bold': True, 'color': builder.COLOR_PRIMARY}),
        ('（离线AI工作室）是一款面向本地/离线环境的', {}),
        ('企业级智能助手平台', {'bold': True}),
        ('，采用"', {}),
        ('大模型决策 + C++ 调度层执行', {'bold': True, 'color': builder.COLOR_ACCENT}),
        ('"的分离式架构，实现 AI 能力的本地私有化部署。', {}),
    ])

    builder.add_heading('1.3 核心设计目标', level=2)
    builder.add_table(
        ['设计目标', '具体措施', '领导视角价值'],
        [
            ['安全可控', 'LLM 只输出 JSON 决策计划，不直接操作系统；所有执行经 C++ 层校验',
             '消除 AI "失控" 风险，满足企业合规审计要求'],
            ['离线优先', '100% 功能可在无网络环境运行，支持本地模型（Ollama 等）',
             '适配内网、涉密、工控等隔离环境'],
            ['透明可追溯', '每一步执行过程在 UI 中可视化展示，计划可人工审查',
             '操作全留痕，符合企业治理规范'],
            ['可扩展', 'Agent 插件化设计，新增能力零侵入核心调度',
             '支持业务定制化扩展，保护长期投资'],
        ],
        col_widths=[3.0, 5.5, 5.5]
    )

    # ==================== 第二章：技术架构总览 ====================
    builder.add_heading('二、技术架构总览', level=1)

    builder.add_heading('2.1 分层架构设计', level=2)
    builder.add_paragraph(
        '本系统采用经典的企业级分层架构设计，自上而下分为四层：UI展示层、调度控制层、Agent执行层和基础设施层。'
        '每一层职责清晰、边界明确，通过 Qt 信号槽机制实现层间通信，达到高内聚低耦合的设计目标。'
    )

    builder.add_paragraph(
        '架构核心特点如下：'
    )
    builder.add_bullet('UI 层与业务逻辑完全分离，仅负责渲染和用户交互')
    builder.add_bullet('调度层作为系统枢纽，协调 LLM 规划与 Agent 执行')
    builder.add_bullet('Agent 层采用插件化设计，可独立扩展新能力')
    builder.add_bullet('基础设施层封装网络、安全、环境检测等通用能力')

    if 1 in D:
        builder.add_image(D[1], width=6.0,
            caption='图 2-1  系统四层架构图',
            description='上图展示了系统的四层架构：最顶层为 UI 展示层（MainWindow、ChatPage 等组件），'
                        '向下依次为调度控制层（Orchestrator、TaskScheduler、Planner）、Agent 执行层（FileAgent、CodeAgent 等五个专业 Agent），'
                        '最底层为基础设施层（LlmClient、SecurityManager、EnvironmentDetector）。'
                        '各层之间通过信号槽松耦合通信，实线表示依赖关系，虚线表示配置或可选依赖。')
    else:
        builder.add_paragraph('【分层架构图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('2.2 架构设计思想', level=2)
    builder.add_paragraph(
        '本系统采用 "决策-执行分离"（Separation of Decision and Execution, SDE）架构模式，'
        '这是区别于传统 Function Calling 方案的核心创新点。在 SDE 模式下：'
    )
    builder.add_bullet('决策层（LLM）：负责理解用户意图，生成结构化的 JSON 执行计划，不接触操作系统')
    builder.add_bullet('调度层（C++）：负责解析计划、校验合法性、按步骤调度 Agent 执行，是安全守门人')
    builder.add_bullet('执行层（Agent）：负责调用操作系统能力完成具体任务，受安全管理器约束')

    builder.add_paragraph('')
    builder.add_rich_paragraph([
        ('这种三层分离设计使系统具备', {}),
        ('银行级的可控性', {'bold': True, 'color': builder.COLOR_ACCENT}),
        ('——即使 AI 产生误判或幻觉，也无法越权执行任何危险操作，因为所有执行必须经过 C++ 层的严格校验。', {}),
    ])

    # ==================== 第三章：核心系统设计 ====================
    builder.add_heading('三、核心系统设计', level=1)

    builder.add_heading('3.1 核心类关系', level=2)
    builder.add_paragraph(
        '系统采用面向对象设计，以 Agent 抽象基类为核心，所有具体 Agent 均继承自该基类。'
        'Orchestrator 作为总控协调器，持有 LlmClient、TaskScheduler、Planner、ConversationManager 等组件的实例，'
        '并通过 addAgent() 方法注册所有可用的 Agent。TaskScheduler 负责将任务步骤路由到正确的 Agent 执行。'
    )

    builder.add_paragraph('核心类的设计遵循以下原则：')
    builder.add_bullet('Agent 基类定义纯虚接口（name、description、tools、executeTool），确保所有 Agent 具有统一行为')
    builder.add_bullet('Tool 类封装工具的元数据（名称、描述、参数列表），支持动态注册和运行时发现')
    builder.add_bullet('SecurityManager 采用静态方法设计，提供全局可用的安全校验能力')
    builder.add_bullet('ThemeManager 采用单例模式，确保全局主题状态一致性')

    if 2 in D:
        builder.add_image(D[2], width=6.5,
            caption='图 3-1  核心类图（UML）',
            description='上图展示了系统核心类的继承和关联关系：Agent 基类定义了所有 Agent 的公共接口；'
                        'FileAgent、ComputerAgent、CodeAgent、SearchAgent、TextAgent 继承自 Agent；'
                        'Orchestrator 也继承自 Agent（历史设计），持有 LlmClient、TaskScheduler、Planner、ConversationManager 的实例；'
                        'TaskScheduler 维护一个 Agent 列表用于工具路由；各 Agent 依赖 SecurityManager 进行安全校验。')
    else:
        builder.add_paragraph('【核心类图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('3.2 模块职责矩阵', level=2)
    builder.add_paragraph(
        '下表详细列出了系统核心模块的类型、职责和采用的设计模式，帮助理解各模块的功能定位和设计理念：'
    )
    builder.add_table(
        ['模块', '类型', '核心职责', '设计模式'],
        [
            ['Orchestrator', '协调器', '两阶段执行总控、信号代理、任务类型识别', '外观模式 + 中介者模式'],
            ['TaskScheduler', '调度器', 'Agent 路由、步骤顺序执行、状态追踪', '责任链模式'],
            ['Planner', '解析器', 'JSON 计划提取、验证、结构化转换', '策略模式'],
            ['PromptBuilder', '构造器', '系统提示词/用户提示词生成', '建造者模式'],
            ['SecurityManager', '安全层', '参数校验、路径安全、命令白名单', '代理模式'],
            ['ConversationManager', '数据层', '多会话管理、历史持久化', '仓库模式'],
            ['ThemeManager', 'UI 基础', '主题状态管理、全局样式生成', '单例模式'],
        ],
        col_widths=[3.5, 2.0, 5.0, 3.5]
    )

    # ==================== 第四章：关键业务流程 ====================
    builder.add_heading('四、关键业务流程', level=1)

    builder.add_heading('4.1 两阶段执行流程', level=2)
    builder.add_paragraph(
        '两阶段执行是本系统最核心的业务流程，体现了"规划与执行分离"的设计哲学。'
        '当用户发送一条任务请求时，系统不会立即执行，而是先进入规划阶段，由大模型生成结构化的 JSON 执行计划，'
        '经过 Planner 解析和校验后，再进入执行阶段，由 TaskScheduler 逐步骤调度 Agent 完成任务。'
    )

    builder.add_paragraph('规划阶段的核心步骤：')
    builder.add_bullet('Orchestrator 调用 PromptBuilder 构造包含所有 Agent 能力的系统提示词')
    builder.add_bullet('LlmClient 发送 HTTP 请求到配置的 LLM API 端点（支持本地模型）')
    builder.add_bullet('LLM 返回 JSON 格式的执行计划，包含目标（goal）和步骤列表（steps）')
    builder.add_bullet('Planner 提取并解析 JSON，生成 TaskPlan 数据结构')

    builder.add_paragraph('执行阶段的核心步骤：')
    builder.add_bullet('TaskScheduler 遍历 TaskPlan 中的每个 TaskStep')
    builder.add_bullet('根据 tool 名称查找拥有该工具的 Agent（工具路由）')
    builder.add_bullet('调用 Agent.executeTool() 执行具体操作，传递参数并收集结果')
    builder.add_bullet('每一步的执行日志通过信号实时推送到 UI 输出面板')

    if 3 in D:
        builder.add_image(D[3], width=6.0,
            caption='图 4-1  两阶段执行时序图',
            description='上图展示了完整的两阶段执行流程时序：用户输入触发 ChatPage 发送消息 → '
                        'Orchestrator 进入规划阶段 → LlmClient 发送请求并接收流式响应 → '
                        'Planner 解析 JSON 计划 → Orchestrator 发射 planGenerated 信号更新 UI → '
                        'TaskScheduler 开始执行阶段 → 逐步骤调用 Agent（FileAgent、CodeAgent 等）→ '
                        '每个步骤的输出通过 stepOutput 信号实时显示。整个过程用户可随时点击停止按钮中断。')
    else:
        builder.add_paragraph('【两阶段执行时序图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('4.2 任务类型识别与路由', level=2)
    builder.add_paragraph(
        '为了优化大模型的规划效率，系统实现了智能的任务类型识别机制。'
        'Orchestrator 在发送请求前，会先分析用户输入和上传的文件列表，识别出任务所属的类型类别，'
        '然后构造针对性的 System Prompt，引导大模型生成更精准的执行计划。'
    )

    builder.add_paragraph('支持的任务类型包括：')
    builder.add_bullet('文件处理类：文档扩写、代码修改、配置更新等，采用"读取→LLM处理→写入"模式')
    builder.add_bullet('代码开发类：项目创建、代码生成、编译构建、代码分析等')
    builder.add_bullet('系统操作类：系统信息获取、进程管理、环境配置等')
    builder.add_bullet('搜索查询类：本地文件搜索、代码符号搜索、文本内容搜索等')
    builder.add_bullet('文本处理类：Base64 编解码、哈希计算、JSON 格式化等')
    builder.add_bullet('知识问答类：纯文本问答，不需要调用任何工具')

    if 4 in D:
        builder.add_image(D[4], width=5.8,
            caption='图 4-2  任务类型识别与路由流程图',
            description='上图展示了任务类型识别的决策流程：首先判断是否包含上传文件，'
                        '有则进入文件处理类任务分支（细分为扩写、修改、修复等）；无文件则通过关键词匹配 '
                        '（如"编译"、"运行"、"系统信息"等）路由到代码开发类、系统操作类、搜索类或文本处理类；'
                        '无法匹配的归为知识问答类，直接让 LLM 生成文本回答。不同类型对应不同的 System Prompt 模板。')
    else:
        builder.add_paragraph('【任务类型识别路由流程图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('4.3 文件处理任务流程', level=2)
    builder.add_paragraph(
        '文件处理任务是本系统支持的重要场景之一，涵盖文档扩写、代码修改、配置更新等多种需求。'
        '该类任务采用特殊的"读取→LLM处理→写入"三步模式，而非通用的规划-执行模式，'
        '目的是让 LLM 能够先读取原始文件内容，再根据用户指令生成完整的新内容。'
    )

    builder.add_paragraph('文件处理流程的特点：')
    builder.add_bullet('FileAgent 先读取原始文件内容（支持多种编码自动检测）')
    builder.add_bullet('构造包含原始内容和用户指令的专用 Prompt，引导 LLM 生成完整内容（而非大纲）')
    builder.add_bullet('LLM 生成的内容经过大纲检测，若仅为大纲则触发补充机制二次请求')
    builder.add_bullet('最终由 FileAgent 写入文件，自动创建不存在的父目录')

    if 5 in D:
        builder.add_image(D[5], width=5.5,
            caption='图 4-3  文件处理任务流程图',
            description='上图展示了文件处理任务的详细流程：用户上传文件 + 编辑指令 → '
                        'Orchestrator 识别为文件处理任务 → FileAgent.readFile 读取原始内容 → '
                        '构建专用 Prompt（原始内容 + 用户指令）→ LLM 生成新内容 → '
                        '判断内容是否仅为大纲，是则触发补充机制 → FileAgent.writeFile 写入文件 → '
                        'UI 显示操作结果。该流程确保文档扩写等任务能生成完整内容。')
    else:
        builder.add_paragraph('【文件处理任务流程图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    # ==================== 第五章：数据模型设计 ====================
    builder.add_heading('五、数据模型设计', level=1)

    builder.add_heading('5.1 核心数据结构', level=2)
    builder.add_paragraph(
        '系统定义了一组精简但完整的数据结构，用于承载 LLM 规划结果和执行状态。'
        '核心结构包括 TaskPlan（执行计划）和 TaskStep（单个执行步骤），以及 ConversationManager 管理的对话历史数据结构。'
        '所有数据结构都设计为可 JSON 序列化，支持持久化存储和跨进程传输。'
    )

    builder.add_paragraph('核心数据结构说明：')
    builder.add_bullet('TaskPlan：包含 planId（唯一标识）、goal（任务目标）、steps（步骤列表）、createdAt（创建时间）')
    builder.add_bullet('TaskStep：包含 stepId（步骤编号）、agent（建议 Agent）、tool（工具名）、args（参数）、status（状态）、output/error（执行结果）')
    builder.add_bullet('Conversation：包含 id、title、messages 列表、isActive 标记，支持多会话切换')
    builder.add_bullet('ToolInfo：环境检测器返回的工具信息，包含 name、path、version、available 字段')

    if 6 in D:
        builder.add_image(D[6], width=5.5,
            caption='图 5-1  核心数据结构类图',
            description='上图展示了核心数据结构的类关系：TaskPlan 包含多个 TaskStep，形成一对多的聚合关系；'
                        'ConversationManager 管理多个 Conversation，每个 Conversation 包含多条 Message；'
                        'EnvironmentDetector 维护一个工具名到 ToolInfo 的映射表。'
                        'TaskStep 内部定义了 StepStatus 枚举（Pending/Running/Completed/Failed/Skipped）表示执行状态。')
    else:
        builder.add_paragraph('【数据结构类图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('5.2 步骤状态机', level=2)
    builder.add_paragraph(
        '每个 TaskStep 在执行过程中会经历多个状态转换，形成一个清晰的状态机。'
        '状态的转换由 TaskScheduler 驱动，并通过信号通知 UI 层实时更新步骤图标和颜色。'
        '状态机的设计确保了执行过程的可追踪性和可调试性。'
    )

    builder.add_paragraph('步骤状态说明：')
    builder.add_bullet('Pending：初始状态，步骤等待执行，显示灰色圆圈图标')
    builder.add_bullet('Running：正在执行，显示紫色旋转图标，此期间可能发射多个 stepOutput 信号')
    builder.add_bullet('Completed：执行成功，显示绿色对勾图标，输出字段填充执行结果')
    builder.add_bullet('Failed：执行失败，显示红色叉号图标，错误信息填充到 error 字段')
    builder.add_bullet('Skipped：被跳过（预留），用于条件分支或用户中断后的后续步骤')

    if 7 in D:
        builder.add_image(D[7], width=4.5,
            caption='图 5-2  步骤状态机图',
            description='上图展示了 TaskStep 的状态转换路径：初始状态为 Pending；'
                        'TaskScheduler 开始执行时状态变为 Running；'
                        '执行成功后状态变为 Completed，执行失败则变为 Failed；'
                        '用户中断时当前步骤继续完成，后续步骤变为 Skipped。'
                        '状态机确保每个步骤有明确的终态，便于统计执行成功率。')
    else:
        builder.add_paragraph('【步骤状态机图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    # ==================== 第六章：接口设计 ====================
    builder.add_heading('六、接口设计', level=1)

    builder.add_heading('6.1 内部模块通信接口', level=2)
    builder.add_paragraph(
        '系统采用 Qt 的信号槽机制作为内部模块通信的标准接口。'
        '这种设计实现了模块间的完全解耦：发送者只需发射信号，无需知道接收者是谁；'
        '接收者只需连接信号，无需知道发送者是谁。Orchestrator 作为信号代理，'
        '将 TaskScheduler 的所有执行信号透传给 MainWindow，由后者更新 UI 组件。'
    )

    builder.add_paragraph('核心信号说明：')
    builder.add_bullet('planGenerated(plan)：规划阶段完成，携带解析后的 TaskPlan')
    builder.add_bullet('stepStarted(stepId)：某步骤开始执行，UI 更新该步骤图标为"运行中"')
    builder.add_bullet('stepOutput(stepId, output)：步骤产生输出，追加到输出面板')
    builder.add_bullet('stepCompleted/stepFailed(stepId, result/error)：步骤执行结果')
    builder.add_bullet('planCompleted(plan)：所有步骤执行完毕')
    builder.add_bullet('logMessage(logType, message)：实时日志输出，类型包括 llm/agent/system/execution')

    if 8 in D:
        builder.add_image(D[8], width=5.5,
            caption='图 6-1  信号槽总线架构图',
            description='上图展示了信号从 Orchestrator 到 UI 的透传路径：左侧为 Orchestrator 发射的 8 个核心信号；'
                        '右侧为 MainWindow 中对应的 8 个槽函数；中间的箭头表示连接关系。'
                        '这种设计使得 TaskScheduler 和 Agent 完全不知道 UI 的存在，实现了业务逻辑与界面展示的彻底分离。')
    else:
        builder.add_paragraph('【信号槽总线架构图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('6.2 Agent 抽象接口', level=2)
    builder.add_paragraph('所有 Agent 遵循统一的抽象契约，确保系统可以一致地发现、调用和管理它们：')
    builder.add_code_block(
        '// Agent 基类纯虚接口\n'
        'virtual QString     name() const = 0;                          // Agent 名称\n'
        'virtual QString     description() const = 0;                   // 功能描述\n'
        'virtual AgentType   type() const = 0;                          // 类型枚举\n'
        'virtual QList<Tool*> tools() const = 0;                        // 工具清单\n'
        'virtual QVariantMap executeTool(const QString& toolName,       // 工具执行入口\n'
        '                                const QVariantMap& args) = 0;',
        'cpp'
    )

    builder.add_heading('6.3 外部 API 接口', level=2)
    builder.add_paragraph('系统通过标准 OpenAI 兼容 API 与 LLM 通信，支持多种本地模型服务：')
    builder.add_table(
        ['参数', '值'],
        [
            ['请求端点', 'POST {configured_api_url}/chat/completions'],
            ['请求格式', '{"model": "qwen2.5:7b", "messages": [...], "stream": true, "temperature": 0.7, "max_tokens": 8192}'],
            ['响应格式', 'Server-Sent Events (SSE) 流式响应'],
            ['兼容模型', 'Ollama / llama.cpp / LM Studio / vLLM / 任何 OpenAI 兼容 API'],
        ],
        col_widths=[3.0, 11.0]
    )

    # ==================== 第七章：状态机设计 ====================
    builder.add_heading('七、状态机设计', level=1)

    builder.add_heading('7.1 Orchestrator 处理状态机', level=2)
    builder.add_paragraph(
        'Orchestrator 作为系统总控，其自身也遵循严格的状态机模型。'
        '状态机确保了 Orchestrator 在任意时刻都处于明确的状态，防止并发请求导致的混乱。'
        '通过 m_isProcessing 和 m_isPlanningPhase 两个标志位，Orchestrator 准确区分规划阶段和执行阶段。'
    )

    builder.add_paragraph('状态转换说明：')
    builder.add_bullet('Idle → Planning：用户发送请求，设置 m_isProcessing=true, m_isPlanningPhase=true')
    builder.add_bullet('Planning → Executing：Planner 成功解析 JSON 计划，设置 m_isPlanningPhase=false')
    builder.add_bullet('Planning → Idle：LLM 返回纯文本回答或解析失败，直接显示消息')
    builder.add_bullet('Executing → Summarizing：TaskScheduler 完成所有步骤，生成执行总结')
    builder.add_bullet('Executing → Idle：用户点击停止按钮，立即中断执行')

    if 9 in D:
        builder.add_image(D[9], width=5.0,
            caption='图 7-1  Orchestrator 处理状态机',
            description='上图展示了 Orchestrator 的状态转换：初始为 Idle 状态；用户发送请求后进入 Planning（规划）状态；'
                        '规划成功后进入 Executing（执行）状态；执行完毕进入 Summarizing（生成总结）状态后回到 Idle；'
                        '用户点击停止时直接回到 Idle。状态机确保 m_isProcessing 标志位正确维护，防止并发请求。')
    else:
        builder.add_paragraph('【Orchestrator 状态机图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('7.2 任务调度器状态机', level=2)
    builder.add_paragraph(
        'TaskScheduler 维护一个独立的状态机，用于管理执行生命周期。'
        '内部还嵌套了一个步骤级别的状态机，用于追踪单个步骤的执行状态。'
        '状态机设计支持随时中断执行，并能在新计划到来时正确重置状态。'
    )

    builder.add_paragraph('状态转换说明：')
    builder.add_bullet('Ready → Running：调用 startExecution()，开始执行计划')
    builder.add_bullet('Running → Completed：所有步骤执行完毕')
    builder.add_bullet('Running → Paused：用户调用 stopExecution()，设置 m_shouldStop=true')
    builder.add_bullet('Paused → Ready：状态重置，等待新计划')
    builder.add_bullet('Completed → Ready：调用 setPlan(newPlan)，开始新的执行周期')

    if 10 in D:
        builder.add_image(D[10], width=5.0,
            caption='图 7-2  TaskScheduler 状态机',
            description='上图展示了 TaskScheduler 的状态转换：外层状态包括 Ready、Running、Paused、Completed；'
                        'Running 状态内部嵌套了步骤级状态机：StepPending → StepRunning → StepCompleted/StepFailed → StepPending（循环）；'
                        '用户调用 stopExecution() 时，当前步骤继续完成，m_shouldStop 标志阻止后续步骤执行。')
    else:
        builder.add_paragraph('【TaskScheduler 状态机图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    # ==================== 第八章：安全架构设计 ====================
    builder.add_heading('八、安全架构设计', level=1)

    builder.add_heading('8.1 纵深防御体系', level=2)
    builder.add_paragraph(
        '安全性是本系统的核心设计目标之一。采用纵深防御（Defense in Depth）策略，'
        '从架构层到系统层共设置六道安全防线，确保即使某一层被绕过，后续层仍能拦截危险操作。'
        '这种设计使系统具备银行级的安全可控性，满足企业合规审计要求。'
    )

    builder.add_paragraph('六层防御机制详解：')
    builder.add_bullet('第一层（架构隔离）：LLM 不直接执行工具，仅输出 JSON 计划，所有执行由 C++ 层完成')
    builder.add_bullet('第二层（参数校验）：检查必填参数是否存在、类型是否正确，防止空指针和类型错误')
    builder.add_bullet('第三层（路径安全）：检测并阻止路径遍历攻击（如 ../../../etc/passwd）')
    builder.add_bullet('第四层（命令白名单）：仅允许执行预定义的安全命令集（如 gcc、cmake、python）')
    builder.add_bullet('第五层（危险指令识别）：拦截格式化、递归删除等高危操作（如 format、del /s）')
    builder.add_bullet('第六层（权限检查）：执行前校验文件系统读写权限，防止越权访问')

    if 11 in D:
        builder.add_image(D[11], width=4.5,
            caption='图 8-1  六层纵深防御体系',
            description='上图展示了纵深防御的六层防线：从左到右依次为架构隔离层（绿色）、参数校验层（蓝色）、'
                        '路径安全层（橙色）、命令白名单层（粉色）、危险指令识别层（紫色）、权限检查层（青色）；'
                        '箭头表示请求逐层穿透，任意一层拦截即返回错误。这种设计确保了即使 LLM 产生幻觉，'
                        '也无法执行任何危害系统的操作。')
    else:
        builder.add_paragraph('【纵深防御体系图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('8.2 安全策略矩阵', level=2)
    builder.add_table(
        ['层级', '安全机制', '防护目标', '实现位置'],
        [
            ['架构层', '决策-执行分离', '防止 AI 幻觉导致危险操作', 'Orchestrator 两阶段设计'],
            ['输入层', '参数必填校验', '防止缺失参数导致异常', 'SecurityManager::validateParams'],
            ['文件层', '路径遍历防护', '防止访问系统敏感目录', 'SecurityManager::validatePath'],
            ['命令层', '白名单机制', '防止执行任意系统命令', 'SecurityManager::validateCommand'],
            ['语义层', '危险指令识别', '拦截格式化、递归删除等', 'SecurityManager::isDangerousCommand'],
            ['系统层', '文件权限校验', '防止越权读写', 'checkRead/WritePermission'],
        ],
        col_widths=[2.0, 3.0, 4.0, 5.0]
    )

    # ==================== 第九章：部署与运行架构 ====================
    builder.add_heading('九、部署与运行架构', level=1)

    builder.add_heading('9.1 部署架构', level=2)
    builder.add_paragraph(
        '本系统采用单机部署模式，无需服务器集群，非常适合内网隔离环境。'
        '应用程序本身是一个独立的可执行文件（OfflineAIStudio.exe），运行时依赖 Qt 动态库。'
        '数据存储采用本地 JSON 文件，位于用户 AppData 目录。可选的本地 LLM 服务（如 Ollama）'
        '需在应用配置中指定 API 端点地址。'
    )

    builder.add_paragraph('部署要点说明：')
    builder.add_bullet('应用程序：单个可执行文件 + Qt 动态库，无需安装，解压即用')
    builder.add_bullet('配置存储：%APPDATA%/OfflineAIStudio/config.json（API 地址、模型名称等）')
    builder.add_bullet('对话历史：%APPDATA%/OfflineAIStudio/conversations.json（多会话历史）')
    builder.add_bullet('本地模型：支持 Ollama、llama.cpp、LM Studio 等服务，通过 HTTP API 通信')
    builder.add_bullet('开发工具链：gcc/g++、cmake、python 等可选，用于 CodeAgent 的编译和运行功能')

    if 12 in D:
        builder.add_image(D[12], width=5.5,
            caption='图 9-1  部署架构图',
            description='上图展示了系统的部署结构：宿主机器上运行 OfflineAIStudio.exe 应用程序；'
                        '本地数据存储包括 config.json（配置）、conversations.json（对话历史）和 QSettings（主题等）；'
                        '可选的本地 LLM 服务（Ollama/llama.cpp/LM Studio）通过 HTTP API 与应用通信；'
                        '可选的开发工具链（gcc/cmake/python）被 CodeAgent 调用用于编译和运行代码。')
    else:
        builder.add_paragraph('【部署架构图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('9.2 运行环境要求', level=2)
    builder.add_table(
        ['组件', '最低要求', '推荐配置'],
        [
            ['操作系统', 'Windows 10 / macOS 12 / Ubuntu 20.04', 'Windows 11 / macOS 14 / Ubuntu 22.04'],
            ['CPU', 'x86_64 双核', 'x86_64 四核以上'],
            ['内存', '4 GB', '16 GB（运行 7B 模型）'],
            ['磁盘', '200 MB', '500 MB'],
            ['网络', '无要求（离线优先）', '内网环境（本地模型）'],
            ['Qt 运行时', 'Qt 6.11+', 'Qt 6.11+'],
        ],
        col_widths=[3.0, 5.5, 5.5]
    )

    # ==================== 第十章：技术选型与约束 ====================
    builder.add_heading('十、技术选型与约束', level=1)

    builder.add_heading('10.1 技术栈选型', level=2)
    builder.add_paragraph(
        '本系统的技术选型遵循"稳定优先、离线友好"的原则。'
        '选择 Qt 6 Widgets 而非 Qt Quick，因为前者更适合桌面生产力工具，API 更稳定；'
        '选择 C++20 而非 Python 或 Go，因为需要与 Qt 深度集成且对性能有一定要求；'
        '选择 JSON 而非 XML 作为计划格式，因为大模型对 JSON 的理解最准确。'
    )

    builder.add_paragraph('核心技术栈：')
    builder.add_bullet('语言标准：C++20，充分利用现代特性（concepts、ranges、format 等）')
    builder.add_bullet('UI 框架：Qt 6.11+ Widgets，成熟稳定、跨平台支持好')
    builder.add_bullet('网络通信：Qt Network + SSE（Server-Sent Events），支持流式响应')
    builder.add_bullet('数据格式：QJsonDocument，原生支持、解析效率高')
    builder.add_bullet('样式方案：QSS（Qt Style Sheets）+ Fusion 风格，跨平台一致外观')
    builder.add_bullet('构建系统：CMake 3.20+，支持多种编译器和 IDE')

    if 13 in D:
        builder.add_image(D[13], width=5.5,
            caption='图 10-1  技术栈选型图',
            description='上图展示了各技术组件的依赖关系：C++20 语言基础 → Qt 6 框架 → '
                        'Qt Network（网络通信）和 Qt Widgets（UI 组件）；网络模块使用 SSE 协议实现流式响应；'
                        'UI 使用 QSS 样式表 + Fusion 风格实现跨平台一致外观；构建使用 CMake。'
                        '所有技术组件均不依赖网络，支持完全离线部署。')
    else:
        builder.add_paragraph('【技术栈选型图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('10.2 选型决策分析', level=2)
    builder.add_table(
        ['技术点', '选型', '备选方案', '决策理由'],
        [
            ['UI 框架', 'Qt Widgets', 'Qt Quick/QML', 'Widgets 更成熟稳定，适合桌面生产力工具'],
            ['通信协议', 'HTTP + SSE', 'WebSocket', 'SSE 更简单可靠，单向流满足需求'],
            ['计划格式', 'JSON', 'XML/YAML', 'LLM 对 JSON 理解最好，解析库成熟'],
            ['架构模式', '决策-执行分离', 'Function Calling', '不依赖特定模型，通用性最强'],
            ['线程模型', '单线程+信号槽', '多线程线程池', 'Qt 事件循环足够，避免并发复杂度'],
        ],
        col_widths=[2.5, 3.0, 3.0, 5.5]
    )

    builder.add_heading('10.3 关键约束', level=2)
    builder.add_paragraph('系统设计遵循以下硬性约束，确保满足离线场景需求：')
    builder.add_bullet('离线优先：所有网络相关功能必须提供离线降级方案，不可因断网而崩溃')
    builder.add_bullet('本地模型兼容：支持 OpenAI 兼容格式的本地模型端点，不依赖特定云服务')
    builder.add_bullet('跨平台：同一套代码支持 Windows/macOS/Linux，无平台特定代码')
    builder.add_bullet('无外部依赖：除 Qt 外不依赖任何第三方库，降低部署复杂度')

    # ==================== 第十一章：扩展性设计 ====================
    builder.add_heading('十一、扩展性设计', level=1)

    builder.add_heading('11.1 Agent 扩展机制', level=2)
    builder.add_paragraph(
        '系统采用插件化架构设计，新增 Agent 无需修改核心调度逻辑。'
        '开发者只需创建一个继承自 Agent 基类的新类，实现纯虚接口，然后在应用启动时调用 Orchestrator.addAgent() '
        '注册即可。系统会自动发现新 Agent 的工具，并将其纳入规划 Prompt 中。'
    )

    builder.add_paragraph('新增 Agent 的标准步骤：')
    builder.add_bullet('继承 Agent 基类，实现 name()、description()、type()、tools()、executeTool() 方法')
    builder.add_bullet('在 initializeTools() 中创建并注册所有工具（Tool 对象）')
    builder.add_bullet('实现每个工具的具体逻辑，返回包含 success/result/error 的 QVariantMap')
    builder.add_bullet('在 main.cpp 或 Orchestrator 初始化处调用 addAgent(new YourAgent()) 注册')
    builder.add_bullet('重启应用，新 Agent 的能力自动出现在 LLM 的规划范围内')

    if 14 in D:
        builder.add_image(D[14], width=4.0,
            caption='图 11-1  Agent 扩展机制流程图',
            description='上图展示了新增 Agent 的标准流程：开发者创建新 Agent 类 → '
                        '继承 Agent 基类 → 实现 name/description/type 方法 → 注册 tools 列表 → '
                        '实现 executeTool 路由逻辑 → 在应用启动时调用 Orchestrator.addAgent() → '
                        '系统自动将新 Agent 纳入能力矩阵，LLM 规划时即可使用。整个过程零侵入核心调度代码。')
    else:
        builder.add_paragraph('【Agent 扩展机制图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    builder.add_heading('11.2 扩展能力矩阵', level=2)
    builder.add_table(
        ['扩展方向', '工作量', '侵入性', '示例场景'],
        [
            ['新增 Agent', '~1 天', '零侵入', '数据库操作 Agent、浏览器控制 Agent'],
            ['新增工具', '~2 小时', '低侵入', 'FileAgent 新增压缩/解压工具'],
            ['新增任务类型', '~4 小时', '低侵入', '支持 "PPT 生成" 任务识别'],
            ['新增主题', '~30 分钟', '零侵入', 'ThemeManager 新增高对比度主题'],
            ['支持新 LLM', '~2 小时', '低侵入', '接入本地 vLLM 服务'],
        ],
        col_widths=[3.0, 2.0, 2.0, 7.0]
    )

    builder.add_heading('11.3 未来架构演进', level=2)
    builder.add_paragraph(
        '当前版本（v2.0）采用顺序执行模型，后续版本规划了多项架构演进，'
        '包括 DAG 并行调度、插件系统、RAG 知识库、代码解释器沙箱等。'
        '这些演进将在保持核心架构不变的前提下，逐步增强系统能力。'
    )

    if 15 in D:
        builder.add_image(D[15], width=5.0,
            caption='图 11-2  版本演进路线图',
            description='上图展示了从 v2.0 到 v5.0 的版本演进路线：当前 v2.0 采用顺序执行模式；'
                        'v3.0 计划支持 DAG 并行调度和插件系统，提升执行效率；'
                        'v4.0 计划引入 RAG 本地知识库，增强文档检索能力；'
                        'v5.0 计划实现代码解释器沙箱，支持隔离的 Python 执行环境。'
                        '每个版本都是前一版本的增量演进，不破坏现有架构。')
    else:
        builder.add_paragraph('【版本演进路线图 - 见源文档 Mermaid 代码】', italic=True, color=RGBColor(0x99, 0x99, 0x99))

    # ==================== 第十二章：项目演进路线 ====================
    builder.add_heading('十二、项目演进路线', level=1)

    builder.add_heading('12.1 版本演进', level=2)
    builder.add_paragraph(
        '本项目经历了从 v1.0 到 v2.0 的重大架构升级，并规划了后续三个版本的功能演进。'
        '每个版本都有明确的核心特性和架构改进目标，确保项目有序发展。'
    )
    builder.add_table(
        ['版本', '核心特性', '架构改进'],
        [
            ['v1.0', 'ReAct 风格工具调用', 'LLM 直接输出工具调用格式'],
            ['v2.0（当前）', 'JSON 计划驱动 + 多 Agent', '决策-执行分离、TaskScheduler 调度'],
            ['v3.0（规划）', 'DAG 并行步骤、插件系统', '步骤依赖图、动态加载 Agent'],
            ['v4.0（规划）', 'RAG 本地知识库', '向量检索、文档 Embedding'],
            ['v5.0（规划）', '代码解释器沙箱', '隔离 Python 执行环境'],
        ],
        col_widths=[3.0, 5.0, 6.0]
    )

    builder.add_heading('12.2 当前版本核心指标', level=2)
    builder.add_paragraph('下表列出了 v2.0 版本的核心量化指标，展示项目的规模和成熟度：')
    builder.add_table(
        ['指标', '数值', '说明'],
        [
            ['代码总行数', '~15,000+ 行', 'C++ 源码（不含注释）'],
            ['核心模块数', '5 个 Agent + 8 个核心类', '高内聚低耦合设计'],
            ['UI 组件数', '12 个自定义组件', '覆盖完整交互场景'],
            ['工具总数', '60+ 个', '覆盖文件/系统/代码/搜索/文本五大领域'],
            ['单元测试覆盖率', '规划中', 'v3.0 目标 80%'],
            ['支持平台', 'Windows/macOS/Linux', 'Qt 跨平台保证'],
        ],
        col_widths=[3.5, 4.5, 6.0]
    )

    # ==================== 附录 ====================
    builder.add_heading('附录', level=1)

    builder.add_heading('A. 术语表', level=2)
    builder.add_table(
        ['术语', '英文', '解释'],
        [
            ['Agent', '智能代理', '系统中负责特定领域任务执行的模块'],
            ['LLM', '大语言模型', '提供自然语言理解与生成能力的 AI 模型'],
            ['SSE', 'Server-Sent Events', '服务器推送事件，用于流式响应'],
            ['Prompt', '提示词', '发送给大模型的输入文本'],
            ['两阶段执行', 'Two-Phase Execution', '先规划后执行的模式'],
            ['QSS', 'Qt Style Sheets', 'Qt 的样式表机制，类似 CSS'],
        ],
        col_widths=[3.0, 4.0, 7.0]
    )

    builder.add_heading('B. 文档变更记录', level=2)
    builder.add_table(
        ['版本', '日期', '变更内容', '作者'],
        [['v1.0', '2026-07', '初始版本，涵盖 v2.0 完整架构', '技术团队']],
        col_widths=[2.0, 3.0, 6.0, 3.0]
    )

    return builder


# ============================================================
# 主流程
# ============================================================

def main():
    print("=" * 60)
    print("OfflineAIStudio 技术设计文档 -> DOCX 转换")
    print("=" * 60)

    # 1. 读取 Markdown
    if not MD_FILE.exists():
        print(f"[ERROR] 找不到源文件: {MD_FILE}")
        sys.exit(1)

    md_text = MD_FILE.read_text(encoding='utf-8')
    print(f"[OK] 读取源文档: {len(md_text)} 字符")

    # 2. 提取 Mermaid 图表
    mermaid_blocks = extract_mermaid_blocks(md_text)
    print(f"[OK] 提取到 {len(mermaid_blocks)} 个 Mermaid 图表")

    # 3. 渲染图表为 PNG
    diagrams = {}
    for i, (dtype, code) in enumerate(mermaid_blocks, 1):
        print(f"\n[{i}/{len(mermaid_blocks)}] 渲染 {dtype} 图表...")
        img_path = render_mermaid_diagram(i, code)
        if img_path:
            diagrams[i] = img_path
            print(f"  [OK] -> {img_path}")
        else:
            print(f"  [FAIL] 图表 {i} 渲染失败，将在文档中标注")

    print(f"\n[OK] 成功渲染 {len(diagrams)}/{len(mermaid_blocks)} 个图表")

    # 4. 构建 DOCX
    print("\n构建 DOCX 文档...")
    builder = build_document(diagrams)

    # 5. 保存
    builder.save(OUTPUT_DOCX)
    print(f"\n[OK] DOCX 已保存到: {OUTPUT_DOCX}")
    print(f"     文件大小: {OUTPUT_DOCX.stat().st_size / 1024:.1f} KB")
    print("\n完成！")


if __name__ == '__main__':
    main()
